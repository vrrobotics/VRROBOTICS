from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x1B, 0x3A, 0x6B)   # headings
MID_BLUE    = RGBColor(0x25, 0x5F, 0xA8)   # sub-headings
ACCENT      = RGBColor(0x2E, 0x86, 0xAB)   # table header bg (set via shading)
LIGHT_GREY  = RGBColor(0xF2, 0xF4, 0xF8)   # alternate table rows
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x1A, 0x1A, 0x1A)
GREEN       = RGBColor(0x15, 0x7A, 0x3C)
ORANGE      = RGBColor(0xD6, 0x6B, 0x00)

def set_cell_bg(cell, hex_color: str):
    """Fill a table cell background."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_page_break(doc):
    doc.add_page_break()

def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    run = h.runs[0] if h.runs else h.add_run(text)
    if level == 1:
        run.font.color.rgb = DARK_BLUE
        run.font.size = Pt(18)
        run.bold = True
    elif level == 2:
        run.font.color.rgb = MID_BLUE
        run.font.size = Pt(14)
        run.bold = True
    elif level == 3:
        run.font.color.rgb = MID_BLUE
        run.font.size = Pt(12)
        run.bold = True
    else:
        run.font.color.rgb = ACCENT
        run.font.size = Pt(11)
        run.bold = True
    return h

def para(doc, text, bold=False, italic=False, color=None, size=10, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(indent * 0.25)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def schema_table(doc, columns, rows):
    """Render a neat schema table: columns = list of (header, width_inches)."""
    ncols = len(columns)
    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0]
    hdr.height = Cm(0.7)
    for i, (col_name, _) in enumerate(columns):
        cell = hdr.cells[i]
        set_cell_bg(cell, "255FA8")         # accent blue
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(col_name)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.add_row()
        bg = "F2F4F8" if ri % 2 == 0 else "FFFFFF"
        for ci, value in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(8.5)
            if ci == 0:            # column name – bold
                run.bold = True

    # Set column widths
    for row in table.rows:
        for ci, (_, w) in enumerate(columns):
            row.cells[ci].width = Inches(w)

    doc.add_paragraph()   # breathing space

def info_box(doc, label, value, label_color=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(9.5)
    if label_color:
        r1.font.color.rgb = label_color
    r2 = p.add_run(value)
    r2.font.size = Pt(9.5)


# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("LMS Platform")
run.bold = True
run.font.size = Pt(32)
run.font.color.rgb = DARK_BLUE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run("Complete Database & Schema Reference")
run2.font.size = Pt(18)
run2.font.color.rgb = MID_BLUE

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.add_run(f"Generated: {datetime.date.today().strftime('%B %d, %Y')}").font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph()
for label, val in [
    ("Architecture",  "Microservices  (Node.js / Sequelize)"),
    ("Cloud DB",      "AWS RDS MySQL  (lucy-devdb.cu1kcwwqaaqb.us-east-1.rds.amazonaws.com : 3306)"),
    ("Total Databases", "3  (lucy_devdb · lms_admin · lms_course)"),
    ("Total Tables",  "45+  core tables  (across all services)"),
    ("ORM",           "Sequelize 6  (JavaScript)"),
    ("Services",      "auth · assessment · college · organization · course · admin · payment · bastion"),
]:
    info_box(doc, label, val, MID_BLUE)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
#  1. OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "1. Architecture Overview", 1)
para(doc,
     "The LMS is built as a microservices platform where each service owns its domain data. "
     "All MySQL databases reside on a single shared AWS RDS instance. Services communicate "
     "over REST; cross-service joins are handled in application code, not via DB-level foreign keys "
     "across databases.", size=10)

doc.add_paragraph()
heading(doc, "1.1  Database Summary", 2)
schema_table(doc,
    [("Database Name", 1.6), ("Service(s)", 2.5), ("Primary Purpose", 3.1)],
    [
        ("lucy_devdb", "auth · assessment · college · organization · admin (read-only)", "Users, roles, assessments, colleges, organisations"),
        ("lms_admin",  "admin-service",                                                  "Courses, sections, lessons, quizzes, certificates, payments"),
        ("lms_course", "course-service",                                                 "Course catalogue and enrolment records"),
    ]
)

heading(doc, "1.2  Service–Port Map", 2)
schema_table(doc,
    [("Service", 2.0), ("Port", 0.8), ("Database", 1.6), ("Role", 2.8)],
    [
        ("bastion-server",      "8000", "—",           "API Gateway / reverse-proxy"),
        ("auth-service",        "8001", "lucy_devdb",  "Authentication & user management"),
        ("course-service",      "8002", "lms_course",  "Course catalogue & enrolments"),
        ("assessment-service",  "8003", "lucy_devdb",  "Pre/post assessments & questions"),
        ("organization-service","8004", "lucy_devdb",  "Organisation management"),
        ("college-service",     "8005", "lucy_devdb",  "College & branch management"),
        ("admin-service",       "4000", "lms_admin",   "Full LMS admin panel"),
        ("payment-service",     "8006", "lms_admin",   "Payments & coupons"),
    ]
)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
#  2. DATABASE: lucy_devdb
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "2. Database: lucy_devdb", 1)
info_box(doc, "Host",    "lucy-devdb.cu1kcwwqaaqb.us-east-1.rds.amazonaws.com : 3306")
info_box(doc, "Engine",  "MySQL 8.x  (AWS RDS)")
info_box(doc, "Services","auth-service · assessment-service · college-service · organization-service · admin-service (read)")
info_box(doc, "Tables",  "users · roles · colleges · branches · organisations · assessments · questionsets · questions · pre_assessment_registrations")
doc.add_paragraph()

# ── 2.1 roles ─────────────────────────────────────────────────────────────────
heading(doc, "2.1  Table: roles", 2)
info_box(doc, "File", r"backend\auth-service\src\db\models\Role.js")
info_box(doc, "Purpose", "Defines the four system roles; every user has exactly one role.")
doc.add_paragraph()
schema_table(doc,
    [("Column", 1.5), ("Data Type", 1.4), ("Constraints", 2.0), ("Description", 2.3)],
    [
        ("roleId", "STRING",                     "PRIMARY KEY",                  "Unique role identifier  (e.g. 'role_001')"),
        ("role",   "ENUM",                        "NOT NULL",                     "Allowed: student · instructor · admin · auditor"),
    ]
)

# ── 2.2 users ─────────────────────────────────────────────────────────────────
heading(doc, "2.2  Table: users", 2)
info_box(doc, "File", r"backend\auth-service\src\db\models\User.js")
info_box(doc, "Purpose", "Central user records for authentication and profile data.")
doc.add_paragraph()
schema_table(doc,
    [("Column", 2.0), ("Data Type", 1.3), ("Constraints", 1.8), ("Description", 2.1)],
    [
        ("userId",            "STRING",   "PRIMARY KEY",           "Unique user ID  (e.g. 'usr_abc123')"),
        ("name",              "STRING",   "NOT NULL",              "Full display name"),
        ("email",             "STRING",   "UNIQUE, NOT NULL",      "Login email address"),
        ("passwordHash",      "STRING",   "NOT NULL",              "Bcrypt-hashed password"),
        ("phone",             "STRING",   "—",                     "Contact phone number"),
        ("dob",               "DATEONLY", "—",                     "Date of birth"),
        ("gender",            "ENUM",     "—",                     "male · female"),
        ("yearOfEducation",   "STRING",   "—",                     "Academic year string"),
        ("branchId",          "STRING",   "—",                     "Reference to branches.branchId"),
        ("collegeId",         "STRING",   "—",                     "Reference to colleges.clgId"),
        ("yearOfStudy",       "INTEGER",  "—",                     "Current year  (1–6)"),
        ("educationLevel",    "ENUM",     "—",                     "inter · bachelor · master · phd · other"),
        ("branch",            "STRING",   "—",                     "Branch name (denormalized)"),
        ("collegeName",       "STRING",   "—",                     "College name (denormalized)"),
        ("graduationYear",    "STRING",   "—",                     "Expected graduation year"),
        ("collegeCode",       "STRING",   "—",                     "College short code"),
        ("orgId",             "STRING",   "—",                     "Reference to organisations.orgId"),
        ("assessmentId",      "STRING",   "—",                     "Reference to assessments.assessmentId"),
        ("programInterested", "STRING",   "—",                     "Preferred programme name"),
        ("profileStatus",     "ENUM",     "DEFAULT 'pending'",     "active · inactive · pending"),
        ("location",          "STRING",   "—",                     "City / region"),
        ("lastLogin",         "DATE",     "—",                     "Timestamp of last successful login"),
        ("preScore",          "INTEGER",  "—",                     "Pre-assessment total score"),
        ("postScore",         "INTEGER",  "—",                     "Post-assessment total score"),
        ("refreshToken",      "STRING",   "MAX 1024",              "JWT refresh token"),
        ("roleId",            "STRING",   "FK → roles, NOT NULL",  "Assigned role"),
        ("createdAt",         "TIMESTAMP","AUTO",                  "Record creation time"),
        ("updatedAt",         "TIMESTAMP","AUTO",                  "Record last-updated time"),
    ]
)

# ── 2.3 colleges ──────────────────────────────────────────────────────────────
heading(doc, "2.3  Table: colleges", 2)
info_box(doc, "File", r"backend\college-service\src\db\models\College.js")
info_box(doc, "Purpose", "Registered colleges linked to organisations.")
doc.add_paragraph()
schema_table(doc,
    [("Column", 1.5), ("Data Type", 1.2), ("Constraints", 1.6), ("Description", 2.9)],
    [
        ("clgId",      "STRING",  "PRIMARY KEY",          "Unique college ID"),
        ("accesskey",  "STRING",  "UNIQUE, NOT NULL",     "Secret access key for college onboarding"),
        ("clgName",    "STRING",  "NOT NULL",             "Official college name"),
        ("clgAddress", "TEXT",    "—",                    "Full postal address"),
        ("orgId",      "STRING",  "FK → organisations",   "Owning organisation"),
        ("branchIds",  "JSON",    "—",                    'Array of branch PKs  e.g. ["CSE","ECE"]'),
        ("createdAt",  "TIMESTAMP","AUTO",                "Record creation time"),
        ("updatedAt",  "TIMESTAMP","AUTO",                "Record last-updated time"),
    ]
)

# ── 2.4 branches ──────────────────────────────────────────────────────────────
heading(doc, "2.4  Table: branches", 2)
info_box(doc, "File", r"backend\college-service\src\db\models\Branch.js")
doc.add_paragraph()
schema_table(doc,
    [("Column", 1.5), ("Data Type", 1.2), ("Constraints", 1.6), ("Description", 2.9)],
    [
        ("branchId",   "STRING", "PRIMARY KEY",       "Short code  e.g. 'CSE', 'ECE'"),
        ("branchName", "STRING", "UNIQUE, NOT NULL",  "Full branch / department name"),
    ]
)

# ── 2.5 organisations ─────────────────────────────────────────────────────────
heading(doc, "2.5  Table: organisations", 2)
info_box(doc, "File", r"backend\organization-service\src\db\models\Organization.js")
doc.add_paragraph()
schema_table(doc,
    [("Column", 1.5), ("Data Type", 1.2), ("Constraints", 1.6), ("Description", 2.9)],
    [
        ("orgId",      "STRING",   "PRIMARY KEY",         "Unique organisation ID"),
        ("accesskey",  "STRING",   "UNIQUE, NOT NULL",    "Onboarding secret key"),
        ("orgName",    "STRING",   "NOT NULL",            "Organisation name"),
        ("orgState",   "STRING",   "—",                   "State"),
        ("orgCountry", "STRING",   "—",                   "Country"),
        ("orgAddress", "TEXT",     "—",                   "Full address"),
        ("orgPin",     "STRING",   "—",                   "PIN / ZIP code"),
        ("createdAt",  "TIMESTAMP","AUTO",                "Creation time"),
        ("updatedAt",  "TIMESTAMP","AUTO",                "Last-updated time"),
    ]
)

add_page_break(doc)

# ── 2.6 assessments ───────────────────────────────────────────────────────────
heading(doc, "2.6  Table: assessments", 2)
info_box(doc, "File", r"backend\assessment-service\src\db\models\Assessment.js")
info_box(doc, "Purpose", "Scheduled pre/post assessment sessions linked to question sets.")
doc.add_paragraph()
schema_table(doc,
    [("Column", 1.6), ("Data Type", 1.2), ("Constraints", 2.0), ("Description", 2.4)],
    [
        ("assessmentId", "STRING",   "PRIMARY KEY",               "Unique assessment ID"),
        ("type",         "ENUM",     "NOT NULL",                  "pre · post"),
        ("setId",        "STRING",   "FK → questionsets, NOT NULL","Linked question set"),
        ("startAt",      "DATE",     "—",                         "Assessment start time"),
        ("endAt",        "DATE",     "—",                         "Assessment end time"),
        ("score",        "FLOAT",    "—",                         "Maximum / achieved score"),
        ("timer",        "INTEGER",  "—",                         "Time limit (seconds)"),
        ("status",       "ENUM",     "DEFAULT 'not-started'",     "not-started · available · in-progress · completed · expired"),
        ("createdAt",    "TIMESTAMP","AUTO",                      "Creation time"),
        ("updatedAt",    "TIMESTAMP","AUTO",                      "Last-updated time"),
    ]
)

# ── 2.7 questionsets ──────────────────────────────────────────────────────────
heading(doc, "2.7  Table: questionsets", 2)
info_box(doc, "File", r"backend\assessment-service\src\db\models\QuestionSet.js")
doc.add_paragraph()
schema_table(doc,
    [("Column", 1.5), ("Data Type", 1.2), ("Constraints", 1.4), ("Description", 3.1)],
    [
        ("setId",     "STRING",   "PRIMARY KEY",  "Unique set ID"),
        ("setName",   "STRING",   "NOT NULL",     "Display name of the question set"),
        ("category",  "STRING",   "—",            "Topic category (e.g. 'Aptitude', 'Coding')"),
        ("questions", "JSON",     "—",            'Ordered array of question IDs  e.g. ["Q1","Q2"]'),
        ("createdAt", "TIMESTAMP","AUTO",         "Creation time"),
        ("updatedAt", "TIMESTAMP","AUTO",         "Last-updated time"),
    ]
)

# ── 2.8 questions (assessment) ────────────────────────────────────────────────
heading(doc, "2.8  Table: questions  (assessment-service)", 2)
info_box(doc, "File", r"backend\assessment-service\src\db\models\Question.js")
doc.add_paragraph()
schema_table(doc,
    [("Column", 1.8), ("Data Type", 1.2), ("Constraints", 1.4), ("Description", 2.8)],
    [
        ("quesId",           "STRING",   "PRIMARY KEY",  "Unique question ID"),
        ("question",         "TEXT",     "NOT NULL",     "Question body text"),
        ("correctAns",       "STRING",   "NOT NULL",     "Key of the correct option (e.g. 'option2')"),
        ("options",          "JSON",     "NOT NULL",     'Key-value map  {"option1":"A","option2":"B",…}'),
        ("category",         "STRING",   "—",            "Question topic category"),
        ("questionSeverity", "ENUM",     "—",            "easy · medium · hard"),
        ("createdAt",        "TIMESTAMP","AUTO",         "Creation time"),
        ("updatedAt",        "TIMESTAMP","AUTO",         "Last-updated time"),
    ]
)

# ── 2.9 pre_assessment_registrations ─────────────────────────────────────────
heading(doc, "2.9  Table: pre_assessment_registrations", 2)
info_box(doc, "File", r"backend\assessment-service\src\db\models\PreAssessmentRegistration.js")
info_box(doc, "Purpose", "Captures walk-in or registered user intake before a pre-assessment.")
doc.add_paragraph()
schema_table(doc,
    [("Column", 2.0), ("Data Type", 1.2), ("Constraints", 1.8), ("Description", 2.2)],
    [
        ("registrationId",      "STRING",  "PRIMARY KEY",        "Unique registration ID"),
        ("userId",              "STRING",  "NULLABLE",           "Auth user ID (null for walk-ins)"),
        ("fullName",            "STRING",  "NOT NULL, MAX 120",  "Candidate full name"),
        ("email",               "STRING",  "NOT NULL, MAX 160",  "Email address (validated)"),
        ("phoneNumber",         "STRING",  "NOT NULL, MAX 20",   "Phone (regex validated)"),
        ("gender",              "ENUM",    "NOT NULL",           "Defined by GENDERS constant"),
        ("selectedProgram",     "ENUM",    "NOT NULL",           "Defined by PROGRAMS constant"),
        ("uploadedCollegeProof","JSON",    "—",                  "File metadata: {fileName, originalName, mimeType, size, url, storedAt}"),
        ("declarationAccepted", "BOOLEAN", "NOT NULL, DEFAULT false","Must be true to submit"),
        ("assessmentStatus",    "ENUM",    "DEFAULT 'registered'","Defined by ASSESSMENT_STATUSES"),
        ("assessmentStartedAt", "DATE",    "—",                  "Timestamp when assessment begun"),
        ("submittedFromIp",     "STRING",  "MAX 64",             "Client IP at submission"),
        ("submittedUserAgent",  "STRING",  "MAX 255",            "Browser user-agent string"),
        ("createdAt",           "TIMESTAMP","AUTO",              "Creation time"),
        ("updatedAt",           "TIMESTAMP","AUTO",              "Last-updated time"),
    ]
)
para(doc, "Indexes: userId · email · selectedProgram · assessmentStatus", italic=True, size=8.5)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
#  3. DATABASE: lms_admin
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "3. Database: lms_admin", 1)
info_box(doc, "Host",    "lucy-devdb.cu1kcwwqaaqb.us-east-1.rds.amazonaws.com : 3306")
info_box(doc, "Engine",  "MySQL 8.x  (AWS RDS)")
info_box(doc, "Service", "admin-service  (port 4000)")
info_box(doc, "Tables",  "users · categories · courses · sections · lessons · quizzes · questions · quiz_submissions · "
               "lesson_completions · lesson_watch_progress · user_progress · pre_assessment_results · "
               "coupons · certificates · live_classes · seo_fields · settings")
doc.add_paragraph()

# ── 3.1 users (admin) ─────────────────────────────────────────────────────────
heading(doc, "3.1  Table: users  (admin-service)", 2)
info_box(doc, "File", r"backend\admin-service\src\models\User.js")
info_box(doc, "Note", "Separate from auth-service users; used by the admin panel for instructor/student profiles.")
doc.add_paragraph()
schema_table(doc,
    [("Column", 1.8), ("Data Type", 1.3), ("Constraints", 1.6), ("Description", 2.5)],
    [
        ("id",               "INTEGER",   "PK, AUTO-INCREMENT",   "Numeric admin user ID"),
        ("role",             "STRING",    "MAX 100",              "Role string  (admin · instructor · student)"),
        ("email",            "STRING",    "UNIQUE, MAX 255",      "Email address"),
        ("status",           "INTEGER",   "—",                    "Account status  (1=active, 0=inactive)"),
        ("name",             "STRING",    "MAX 255",              "Display name"),
        ("phone",            "STRING",    "MAX 255",              "Contact phone"),
        ("website",          "STRING",    "MAX 255",              "Personal/company website"),
        ("skills",           "TEXT",      "—",                    "Comma-separated skill list"),
        ("facebook",         "TEXT",      "—",                    "Facebook profile URL"),
        ("twitter",          "STRING",    "MAX 255",              "Twitter handle / URL"),
        ("linkedin",         "STRING",    "MAX 255",              "LinkedIn URL"),
        ("address",          "STRING",    "MAX 255",              "Physical address"),
        ("college_name",     "STRING",    "MAX 255",              "College name (denormalized)"),
        ("college_id",       "STRING",    "NULLABLE, MAX 255",    "FK → lucy_devdb.colleges.clgId"),
        ("about",            "TEXT",      "—",                    "Short bio"),
        ("biography",        "LONGTEXT",  "—",                    "Full biography"),
        ("educations",       "LONGTEXT",  "—",                    "Education history (JSON or text)"),
        ("photo",            "STRING",    "MAX 255",              "Profile photo path/URL"),
        ("email_verified_at","DATE",      "—",                    "Email verification timestamp"),
        ("password",         "STRING",    "MAX 255",              "Hashed password"),
        ("remember_token",   "STRING",    "MAX 100",              "Remember-me token"),
        ("paymentkeys",      "LONGTEXT",  "—",                    "Stored payment gateway keys (JSON)"),
        ("video_url",        "STRING",    "MAX 255",              "Intro video URL"),
        ("created_at",       "TIMESTAMP", "AUTO",                 "Creation time"),
        ("updated_at",       "TIMESTAMP", "AUTO",                 "Last-updated time"),
    ]
)

# ── 3.2 categories ────────────────────────────────────────────────────────────
heading(doc, "3.2  Table: categories", 2)
info_box(doc, "File", r"backend\admin-service\src\models\Category.js")
doc.add_paragraph()
schema_table(doc,
    [("Column", 1.6), ("Data Type", 1.3), ("Constraints", 1.5), ("Description", 2.8)],
    [
        ("id",            "INTEGER",   "PK, AUTO-INCREMENT",    "Category ID"),
        ("parent_id",     "INTEGER",   "NULLABLE",              "Self-ref FK for sub-categories"),
        ("title",         "STRING",    "NOT NULL, MAX 255",     "Category name"),
        ("slug",          "STRING",    "MAX 255",               "URL-friendly slug"),
        ("icon",          "STRING",    "MAX 255",               "Icon class or path"),
        ("sort",          "INTEGER",   "DEFAULT 0",             "Display order"),
        ("status",        "STRING",    "DEFAULT 'active'",      "active · inactive"),
        ("keywords",      "STRING",    "MAX 400",               "SEO keywords"),
        ("description",   "STRING",    "MAX 500",               "Short description"),
        ("thumbnail",     "STRING",    "MAX 255",               "Thumbnail image path"),
        ("category_logo", "STRING",    "MAX 255",               "Logo image path"),
        ("created_at",    "TIMESTAMP", "AUTO",                  "Creation time"),
        ("updated_at",    "TIMESTAMP", "AUTO",                  "Last-updated time"),
    ]
)

# ── 3.3 courses (admin) ───────────────────────────────────────────────────────
heading(doc, "3.3  Table: courses  (admin-service)", 2)
info_box(doc, "File", r"backend\admin-service\src\models\Course.js")
info_box(doc, "Purpose", "Full course metadata managed by the admin panel.")
doc.add_paragraph()
schema_table(doc,
    [("Column", 1.9), ("Data Type", 1.3), ("Constraints", 1.5), ("Description", 2.5)],
    [
        ("id",                    "INTEGER", "PK, AUTO-INCREMENT",  "Course ID"),
        ("title",                 "STRING",  "NOT NULL, MAX 255",   "Course title"),
        ("slug",                  "STRING",  "MAX 255",             "URL slug"),
        ("short_description",     "TEXT",    "—",                   "Brief summary"),
        ("user_id",               "INTEGER", "FK → users",          "Creator / instructor ID"),
        ("category_id",           "INTEGER", "FK → categories",     "Course category"),
        ("course_type",           "STRING",  "DEFAULT 'general'",   "general · live · bootcamp"),
        ("status",                "STRING",  "DEFAULT 'active'",    "active · draft · pending · inactive"),
        ("level",                 "STRING",  "MAX 50",              "beginner · intermediate · advanced"),
        ("language",              "STRING",  "MAX 50",              "Course language"),
        ("is_paid",               "BOOLEAN", "DEFAULT false",       "Paid course flag"),
        ("is_best",               "BOOLEAN", "DEFAULT false",       "Featured / bestseller flag"),
        ("price",                 "FLOAT",   "DEFAULT 0",           "Regular price"),
        ("discounted_price",      "FLOAT",   "DEFAULT 0",           "Sale price"),
        ("discount_flag",         "BOOLEAN", "DEFAULT false",       "Discount active flag"),
        ("enable_drip_content",   "BOOLEAN", "DEFAULT false",       "Drip scheduling enabled"),
        ("drip_content_settings", "TEXT",    "—",                   "Drip schedule JSON config"),
        ("meta_keywords",         "TEXT",    "—",                   "SEO meta keywords"),
        ("meta_description",      "TEXT",    "—",                   "SEO meta description"),
        ("thumbnail",             "STRING",  "MAX 255",             "Thumbnail image path"),
        ("banner",                "STRING",  "MAX 255",             "Banner image path"),
        ("preview",               "STRING",  "MAX 255",             "Preview video path/URL"),
        ("description",           "LONGTEXT","—",                   "Full HTML description"),
        ("requirements",          "TEXT",    "—",                   "Pre-requisites (JSON list)"),
        ("outcomes",              "TEXT",    "—",                   "Learning outcomes (JSON list)"),
        ("faqs",                  "TEXT",    "—",                   "FAQ items (JSON)"),
        ("instructor_ids",        "TEXT",    "—",                   "Co-instructor IDs (JSON/CSV)"),
        ("average_rating",        "FLOAT",   "DEFAULT 0",           "Computed average rating"),
        ("expiry_period",         "INTEGER", "NULLABLE",            "Access validity in days"),
        ("created_at",            "TIMESTAMP","AUTO",               "Creation time"),
        ("updated_at",            "TIMESTAMP","AUTO",               "Last-updated time"),
    ]
)

add_page_break(doc)

# ── 3.4 sections ──────────────────────────────────────────────────────────────
heading(doc, "3.4  Table: sections", 2)
info_box(doc, "File", r"backend\admin-service\src\models\Section.js")
doc.add_paragraph()
schema_table(doc,
    [("Column", 1.5), ("Data Type", 1.3), ("Constraints", 1.5), ("Description", 2.9)],
    [
        ("id",         "INTEGER",   "PK, AUTO-INCREMENT", "Section ID"),
        ("user_id",    "INTEGER",   "—",                  "Creator user ID"),
        ("course_id",  "INTEGER",   "FK → courses",       "Parent course"),
        ("title",      "STRING",    "MAX 255",            "Section heading"),
        ("sort",       "INTEGER",   "—",                  "Display order within course"),
        ("created_at", "TIMESTAMP", "AUTO",               "Creation time"),
        ("updated_at", "TIMESTAMP", "AUTO",               "Last-updated time"),
    ]
)

# ── 3.5 lessons ───────────────────────────────────────────────────────────────
heading(doc, "3.5  Table: lessons", 2)
info_box(doc, "File", r"backend\admin-service\src\models\Lesson.js")
doc.add_paragraph()
schema_table(doc,
    [("Column", 1.7), ("Data Type", 1.3), ("Constraints", 1.4), ("Description", 2.8)],
    [
        ("id",              "INTEGER",   "PK, AUTO-INCREMENT", "Lesson ID"),
        ("title",           "STRING",    "MAX 255",            "Lesson title"),
        ("user_id",         "INTEGER",   "—",                  "Author user ID"),
        ("course_id",       "INTEGER",   "FK → courses",       "Parent course"),
        ("section_id",      "INTEGER",   "FK → sections",      "Parent section"),
        ("lesson_type",     "STRING",    "MAX 255",            "video · text · quiz · live"),
        ("duration",        "STRING",    "MAX 255",            'Duration string  e.g. "45 min"'),
        ("total_mark",      "INTEGER",   "—",                  "Max marks (quiz-type lessons)"),
        ("pass_mark",       "INTEGER",   "—",                  "Passing marks threshold"),
        ("retake",          "INTEGER",   "—",                  "Number of allowed retakes"),
        ("lesson_src",      "STRING",    "MAX 255",            "Video URL or source path"),
        ("attachment",      "LONGTEXT",  "—",                  "Attached file data (JSON array)"),
        ("attachment_type", "STRING",    "MAX 255",            "File MIME type"),
        ("video_type",      "TEXT",      "—",                  "youtube · vimeo · uploaded"),
        ("thumbnail",       "STRING",    "MAX 255",            "Lesson thumbnail"),
        ("is_free",         "INTEGER",   "—",                  "1 = free preview lesson"),
        ("sort",            "INTEGER",   "—",                  "Display order within section"),
        ("description",     "LONGTEXT",  "—",                  "Lesson body / transcript"),
        ("summary",         "LONGTEXT",  "—",                  "Summary notes"),
        ("status",          "INTEGER",   "—",                  "1=published, 0=draft"),
        ("created_at",      "TIMESTAMP", "AUTO",               "Creation time"),
        ("updated_at",      "TIMESTAMP", "AUTO",               "Last-updated time"),
    ]
)

# ── 3.6 quizzes ───────────────────────────────────────────────────────────────
heading(doc, "3.6  Table: quizzes", 2)
info_box(doc, "File", r"backend\admin-service\src\models\Quiz.js")
doc.add_paragraph()
schema_table(doc,
    [("Column", 1.6), ("Data Type", 1.3), ("Constraints", 1.5), ("Description", 2.8)],
    [
        ("id",          "BIGINT UNSIGNED","PK, AUTO-INCREMENT","Quiz ID"),
        ("course_id",   "BIGINT UNSIGNED","FK → courses",      "Parent course"),
        ("section_id",  "BIGINT UNSIGNED","FK → sections",     "Parent section"),
        ("title",       "STRING",         "NOT NULL, MAX 255", "Quiz title"),
        ("slug",        "STRING",         "MAX 255",           "URL slug"),
        ("description", "TEXT",           "—",                 "Quiz instructions"),
        ("total_marks", "INTEGER",        "DEFAULT 0",         "Sum of all question marks"),
        ("pass_marks",  "INTEGER",        "DEFAULT 0",         "Minimum passing score"),
        ("duration",    "INTEGER",        "—",                 "Time limit in minutes"),
        ("serial",      "INTEGER",        "DEFAULT 0",         "Display order"),
        ("status",      "INTEGER",        "DEFAULT 1",         "1=active, 0=inactive"),
        ("created_at",  "TIMESTAMP",      "AUTO",              "Creation time"),
        ("updated_at",  "TIMESTAMP",      "AUTO",              "Last-updated time"),
    ]
)

# ── 3.7 questions (quiz) ──────────────────────────────────────────────────────
heading(doc, "3.7  Table: questions  (admin-service / quiz)", 2)
info_box(doc, "File", r"backend\admin-service\src\models\Question.js")
doc.add_paragraph()
schema_table(doc,
    [("Column", 1.5), ("Data Type", 1.3), ("Constraints", 1.5), ("Description", 2.9)],
    [
        ("id",         "INTEGER",   "PK, AUTO-INCREMENT", "Question ID"),
        ("quiz_id",    "INTEGER",   "FK → quizzes",       "Parent quiz"),
        ("title",      "LONGTEXT",  "—",                  "Question text / body"),
        ("type",       "STRING",    "MAX 255",            "multiple-choice · short-answer · essay"),
        ("answer",     "MEDIUMTEXT","—",                  "Correct answer / answer key"),
        ("options",    "LONGTEXT",  "—",                  "Answer options (JSON array)"),
        ("sort",       "INTEGER",   "—",                  "Display order in quiz"),
        ("created_at", "TIMESTAMP", "AUTO",               "Creation time"),
        ("updated_at", "TIMESTAMP", "AUTO",               "Last-updated time"),
    ]
)

add_page_break(doc)

# ── 3.8 quiz_submissions ──────────────────────────────────────────────────────
heading(doc, "3.8  Table: quiz_submissions", 2)
info_box(doc, "File", r"backend\admin-service\src\models\QuizSubmission.js")
doc.add_paragraph()
schema_table(doc,
    [("Column", 1.7), ("Data Type", 1.3), ("Constraints", 1.4), ("Description", 2.8)],
    [
        ("id",             "INTEGER",  "PK, AUTO-INCREMENT","Submission ID"),
        ("quiz_id",        "INTEGER",  "FK → quizzes",      "Quiz attempted"),
        ("user_id",        "INTEGER",  "FK → users",        "Student who submitted"),
        ("correct_answer", "LONGTEXT", "—",                 "JSON: correctly answered questions"),
        ("wrong_answer",   "LONGTEXT", "—",                 "JSON: incorrectly answered questions"),
        ("submits",        "LONGTEXT", "—",                 "JSON array of all submission attempts"),
        ("created_at",     "TIMESTAMP","AUTO",              "Creation time"),
        ("updated_at",     "TIMESTAMP","AUTO",              "Last-updated time"),
    ]
)

# ── 3.9 lesson_completions ────────────────────────────────────────────────────
heading(doc, "3.9  Table: lesson_completions", 2)
info_box(doc, "File", r"backend\admin-service\src\models\LessonCompletion.js")
info_box(doc, "Purpose", "One row per (user, lesson) inserted once the lesson is completed. Drives sidebar tick marks and course progress %.")
doc.add_paragraph()
schema_table(doc,
    [("Column", 1.5), ("Data Type", 1.3), ("Constraints", 1.8), ("Description", 2.6)],
    [
        ("id",         "INTEGER",  "PK, AUTO-INCREMENT",      "Row ID"),
        ("user_id",    "BIGINT",   "NOT NULL",                "Auth-service user ID"),
        ("course_id",  "INTEGER",  "NOT NULL",                "Course the lesson belongs to"),
        ("lesson_id",  "INTEGER",  "NOT NULL",                "Completed lesson"),
        ("created_at", "TIMESTAMP","AUTO",                    "Completion timestamp"),
        ("updated_at", "TIMESTAMP","AUTO",                    "Last-updated time"),
    ]
)
para(doc, "Unique index: (user_id, lesson_id)   |   Index: (user_id, course_id)", italic=True, size=8.5)

# ── 3.10 lesson_watch_progress ────────────────────────────────────────────────
heading(doc, "3.10  Table: lesson_watch_progress", 2)
info_box(doc, "File", r"backend\admin-service\src\models\LessonWatchProgress.js")
info_box(doc, "Purpose", "Tracks the highest second reached in a video; enables resume and 30% completion auto-trigger.")
doc.add_paragraph()
schema_table(doc,
    [("Column", 1.7), ("Data Type", 1.3), ("Constraints", 1.8), ("Description", 2.4)],
    [
        ("id",               "INTEGER",  "PK, AUTO-INCREMENT", "Row ID"),
        ("user_id",          "BIGINT",   "NOT NULL",           "Auth-service user ID"),
        ("course_id",        "INTEGER",  "NOT NULL",           "Parent course"),
        ("lesson_id",        "INTEGER",  "NOT NULL",           "Video lesson"),
        ("current_duration", "INTEGER",  "NOT NULL, DEFAULT 0","Highest second watched"),
        ("created_at",       "TIMESTAMP","AUTO",               "First watch timestamp"),
        ("updated_at",       "TIMESTAMP","AUTO",               "Last-updated time"),
    ]
)
para(doc, "Unique index: (user_id, lesson_id)   |   Index: (user_id, course_id)", italic=True, size=8.5)

# ── 3.11 user_progress ────────────────────────────────────────────────────────
heading(doc, "3.11  Table: user_progress", 2)
info_box(doc, "File", r"backend\admin-service\src\models\UserProgress.js")
info_box(doc, "Purpose", "One row per (user, program) — tracks overall enrolment and last-watched lesson.")
doc.add_paragraph()
schema_table(doc,
    [("Column", 1.6), ("Data Type", 1.3), ("Constraints", 1.8), ("Description", 2.5)],
    [
        ("id",             "INTEGER",  "PK, AUTO-INCREMENT", "Row ID"),
        ("user_id",        "BIGINT",   "NOT NULL",           "Auth-service user ID"),
        ("program_id",     "INTEGER",  "NOT NULL",           "Programme enrolled in"),
        ("course_id",      "INTEGER",  "NULLABLE",           "Current course being studied"),
        ("last_lesson_id", "INTEGER",  "NULLABLE",           "Last lesson watched"),
        ("enrolled",       "BOOLEAN",  "NOT NULL, DEFAULT false","True once enrolment confirmed"),
        ("created_at",     "TIMESTAMP","AUTO",               "Enrolment timestamp"),
        ("updated_at",     "TIMESTAMP","AUTO",               "Last-updated time"),
    ]
)
para(doc, "Unique index: (user_id, program_id)   |   Index: (user_id)", italic=True, size=8.5)

add_page_break(doc)

# ── 3.12 pre_assessment_results ───────────────────────────────────────────────
heading(doc, "3.12  Table: pre_assessment_results", 2)
info_box(doc, "File", r"backend\admin-service\src\models\PreAssessmentResult.js")
doc.add_paragraph()
schema_table(doc,
    [("Column", 1.6), ("Data Type", 1.3), ("Constraints", 1.5), ("Description", 2.8)],
    [
        ("id",         "INTEGER",  "PK, AUTO-INCREMENT","Row ID"),
        ("user_id",    "BIGINT",   "NOT NULL",          "Auth-service user ID"),
        ("program_id", "INTEGER",  "NULLABLE",          "Programme assessed for"),
        ("score",      "FLOAT",    "NOT NULL, DEFAULT 0","Achieved score"),
        ("passed",     "BOOLEAN",  "NOT NULL, DEFAULT false","True if score ≥ passing threshold"),
        ("created_at", "TIMESTAMP","AUTO",              "Result recorded time"),
        ("updated_at", "TIMESTAMP","AUTO",              "Last-updated time"),
    ]
)
para(doc, "Index: (user_id, program_id)", italic=True, size=8.5)

# ── 3.13 coupons ──────────────────────────────────────────────────────────────
heading(doc, "3.13  Table: coupons", 2)
info_box(doc, "File", r"backend\admin-service\src\models\Coupon.js")
doc.add_paragraph()
schema_table(doc,
    [("Column", 1.5), ("Data Type", 1.3), ("Constraints", 1.5), ("Description", 2.9)],
    [
        ("id",         "INTEGER",  "PK, AUTO-INCREMENT","Coupon ID"),
        ("user_id",    "INTEGER",  "NOT NULL, FK → users","Creator / owner"),
        ("code",       "STRING",   "UNIQUE, NOT NULL",  "Coupon code string"),
        ("discount",   "INTEGER",  "NOT NULL",          "Discount value (% or fixed amount)"),
        ("expiry",     "BIGINT",   "NOT NULL",          "Expiry as Unix timestamp"),
        ("status",     "TINYINT",  "DEFAULT 1",         "1=active, 0=inactive"),
        ("created_at", "TIMESTAMP","AUTO",              "Creation time"),
        ("updated_at", "TIMESTAMP","AUTO",              "Last-updated time"),
    ]
)

# ── 3.14 certificates ─────────────────────────────────────────────────────────
heading(doc, "3.14  Table: certificates", 2)
info_box(doc, "File", r"backend\admin-service\src\models\Certificate.js")
info_box(doc, "Purpose", "Issued certificates; identifier is used for public download URL.")
doc.add_paragraph()
schema_table(doc,
    [("Column", 1.7), ("Data Type", 1.3), ("Constraints", 1.5), ("Description", 2.7)],
    [
        ("id",              "INTEGER",  "PK, AUTO-INCREMENT",   "Certificate ID"),
        ("user_id",         "STRING",   "NULLABLE, MAX 255",    "Auth-service userId (string format)"),
        ("course_id",       "INTEGER",  "NULLABLE",             "Course the certificate is for"),
        ("identifier",      "STRING",   "UNIQUE, MAX 100",      "Public download identifier token"),
        ("title",           "STRING",   "NULLABLE, MAX 255",    "Certificate title"),
        ("description",     "TEXT",     "NULLABLE",             "Certificate description"),
        ("template_image",  "STRING",   "NULLABLE, MAX 255",    "Template image path"),
        ("status",          "TINYINT",  "DEFAULT 1",            "1=visible, 0=hidden"),
        ("issued_at",       "DATE",     "NULLABLE",             "Issue date"),
        ("created_at",      "TIMESTAMP","AUTO",                 "Creation time"),
        ("updated_at",      "TIMESTAMP","AUTO",                 "Last-updated time"),
    ]
)

# ── 3.15 live_classes ─────────────────────────────────────────────────────────
heading(doc, "3.15  Table: live_classes", 2)
info_box(doc, "File", r"backend\admin-service\src\models\LiveClass.js")
doc.add_paragraph()
schema_table(doc,
    [("Column", 1.9), ("Data Type", 1.3), ("Constraints", 1.3), ("Description", 2.7)],
    [
        ("id",                    "INTEGER",  "PK, AUTO-INCREMENT","Session ID"),
        ("user_id",               "INTEGER",  "—",                 "Instructor user ID"),
        ("course_id",             "INTEGER",  "—",                 "Course this session belongs to"),
        ("class_topic",           "STRING",   "MAX 255",           "Session topic / title"),
        ("provider",              "STRING",   "MAX 255",           "zoom · jitsi · bigbluebutton"),
        ("class_date_and_time",   "DATE",     "—",                 "Scheduled date & time"),
        ("additional_info",       "LONGTEXT", "—",                 "Extra session info (JSON or text)"),
        ("note",                  "TEXT",     "—",                 "Instructor notes"),
        ("created_at",            "TIMESTAMP","AUTO",              "Creation time"),
        ("updated_at",            "TIMESTAMP","AUTO",              "Last-updated time"),
    ]
)

add_page_break(doc)

# ── 3.16 seo_fields ───────────────────────────────────────────────────────────
heading(doc, "3.16  Table: seo_fields", 2)
info_box(doc, "File", r"backend\admin-service\src\models\SeoField.js")
doc.add_paragraph()
schema_table(doc,
    [("Column", 1.8), ("Data Type", 1.2), ("Constraints", 1.3), ("Description", 2.9)],
    [
        ("id",               "INTEGER",  "PK, AUTO-INCREMENT","Row ID"),
        ("course_id",        "INTEGER",  "—",                 "Associated course (nullable for pages)"),
        ("route",            "STRING",   "MAX 255",           "URL route pattern"),
        ("name_route",       "STRING",   "MAX 255",           "Named route identifier"),
        ("meta_title",       "STRING",   "MAX 255",           "HTML <title> tag content"),
        ("meta_description", "TEXT",     "—",                 "Meta description content"),
        ("meta_keywords",    "TEXT",     "—",                 "Meta keywords"),
        ("meta_robot",       "STRING",   "MAX 100",           "Robots directive  (index, follow…)"),
        ("canonical_url",    "STRING",   "MAX 255",           "Canonical URL"),
        ("custom_url",       "STRING",   "MAX 255",           "Custom redirect / alias URL"),
        ("json_ld",          "TEXT",     "—",                 "Structured data JSON-LD script"),
        ("og_title",         "STRING",   "MAX 255",           "Open Graph title"),
        ("og_description",   "TEXT",     "—",                 "Open Graph description"),
        ("og_image",         "STRING",   "MAX 255",           "Open Graph image URL"),
        ("created_at",       "TIMESTAMP","AUTO",              "Creation time"),
        ("updated_at",       "TIMESTAMP","AUTO",              "Last-updated time"),
    ]
)

# ── 3.17 settings ─────────────────────────────────────────────────────────────
heading(doc, "3.17  Table: settings", 2)
info_box(doc, "File", r"backend\admin-service\src\models\Setting.js")
info_box(doc, "Purpose", "Key-value store for system-wide configuration (certificate templates, payment keys, etc.).")
doc.add_paragraph()
schema_table(doc,
    [("Column", 1.5), ("Data Type", 1.3), ("Constraints", 1.3), ("Description", 3.1)],
    [
        ("id",          "INTEGER",  "PK, AUTO-INCREMENT","Setting ID"),
        ("type",        "STRING",   "MAX 255",           "Setting key / type identifier"),
        ("description", "LONGTEXT", "—",                 "Setting value (JSON blob or plain text)"),
        ("created_at",  "TIMESTAMP","AUTO",              "Creation time"),
        ("updated_at",  "TIMESTAMP","AUTO",              "Last-updated time"),
    ]
)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
#  4. DATABASE: lms_course
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "4. Database: lms_course", 1)
info_box(doc, "Host",    "lucy-devdb.cu1kcwwqaaqb.us-east-1.rds.amazonaws.com : 3306")
info_box(doc, "Engine",  "MySQL 8.x  (AWS RDS)")
info_box(doc, "Service", "course-service  (port 8002)")
info_box(doc, "Tables",  "courses · enrollments")
doc.add_paragraph()

# ── 4.1 courses (course-service) ──────────────────────────────────────────────
heading(doc, "4.1  Table: courses  (course-service)", 2)
info_box(doc, "File", r"backend\course-service\src\db\models\Course.js")
info_box(doc, "Note", "Lightweight course catalogue separate from the admin-service courses; holds domain-level course identity and module metadata.")
doc.add_paragraph()
schema_table(doc,
    [("Column", 2.0), ("Data Type", 1.3), ("Constraints", 1.4), ("Description", 2.5)],
    [
        ("courseId",               "STRING",   "PRIMARY KEY",        "Domain course ID  (e.g. 'CRS_001')"),
        ("title",                  "STRING",   "NOT NULL",           "Course title"),
        ("description",            "TEXT",     "—",                  "Course description"),
        ("duration",               "INTEGER",  "NOT NULL",           "Total duration (hours)"),
        ("isPreAssessmentNeeded",  "BOOLEAN",  "DEFAULT false",      "Requires pre-assessment before enrolment"),
        ("modules",                "JSON",     "—",                  'Array of module objects  [{moduleId, title, duration}]'),
        ("createdAt",              "TIMESTAMP","AUTO",               "Creation time"),
        ("updatedAt",              "TIMESTAMP","AUTO",               "Last-updated time"),
    ]
)

# ── 4.2 enrollments ───────────────────────────────────────────────────────────
heading(doc, "4.2  Table: enrollments", 2)
info_box(doc, "File", r"backend\course-service\src\db\models\Enrollment.js")
info_box(doc, "Purpose", "Student enrolment records; userId is a logical FK to auth-service (no DB-level constraint across databases).")
doc.add_paragraph()
schema_table(doc,
    [("Column", 1.7), ("Data Type", 1.3), ("Constraints", 1.8), ("Description", 2.4)],
    [
        ("enrollmentId", "STRING",   "PRIMARY KEY",               "Unique enrolment ID"),
        ("userId",       "STRING",   "NOT NULL",                  "Auth-service userId (cross-DB logical FK)"),
        ("courseId",     "STRING",   "NOT NULL, FK → courses",    "Enrolled course"),
        ("status",       "ENUM",     "DEFAULT 'enrolled'",        "enrolled · in-progress · completed · dropped"),
        ("enrolledAt",   "DATE",     "DEFAULT NOW",               "Enrolment date"),
        ("completedAt",  "DATE",     "NULLABLE",                  "Completion date (null if not done)"),
        ("createdAt",    "TIMESTAMP","AUTO",                      "Record creation time"),
        ("updatedAt",    "TIMESTAMP","AUTO",                      "Record last-updated time"),
    ]
)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
#  5. CROSS-DATABASE RELATIONSHIPS
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "5. Cross-Database Relationships", 1)
para(doc,
     "Because all tables live on the same AWS RDS instance but in different databases, "
     "foreign keys that cross database boundaries are enforced at the application layer (Sequelize / Express), "
     "not at the MySQL level. The table below documents these logical relationships.", size=10)
doc.add_paragraph()
schema_table(doc,
    [("Source (DB.Table.Column)", 2.2), ("→ Target (DB.Table.Column)", 2.2), ("Enforcement", 1.4), ("Notes", 1.4)],
    [
        ("lms_admin.users.college_id",       "lucy_devdb.colleges.clgId",             "App layer",  "Admin user linked to college"),
        ("lms_course.enrollments.userId",    "lucy_devdb.users.userId",               "App layer",  "Cross-DB enrolment → auth user"),
        ("lucy_devdb.users.roleId",          "lucy_devdb.roles.roleId",               "DB FK",      "Same-DB FK enforced by MySQL"),
        ("lucy_devdb.users.collegeId",       "lucy_devdb.colleges.clgId",             "App layer",  "Denormalized reference"),
        ("lucy_devdb.users.orgId",           "lucy_devdb.organisations.orgId",        "App layer",  "User's organisation"),
        ("lucy_devdb.users.assessmentId",    "lucy_devdb.assessments.assessmentId",   "App layer",  "Assigned assessment"),
        ("lucy_devdb.colleges.orgId",        "lucy_devdb.organisations.orgId",        "App layer",  "College → Organisation"),
        ("lucy_devdb.assessments.setId",     "lucy_devdb.questionsets.setId",         "DB FK",      "Assessment question set"),
        ("lms_admin.courses.user_id",        "lms_admin.users.id",                    "App layer",  "Course creator"),
        ("lms_admin.courses.category_id",    "lms_admin.categories.id",               "App layer",  "Course category"),
        ("lms_admin.sections.course_id",     "lms_admin.courses.id",                  "App layer",  "Section → Course"),
        ("lms_admin.lessons.course_id",      "lms_admin.courses.id",                  "App layer",  "Lesson → Course"),
        ("lms_admin.lessons.section_id",     "lms_admin.sections.id",                 "App layer",  "Lesson → Section"),
        ("lms_admin.quizzes.course_id",      "lms_admin.courses.id",                  "App layer",  "Quiz → Course"),
        ("lms_admin.questions.quiz_id",      "lms_admin.quizzes.id",                  "App layer",  "Question → Quiz"),
        ("lms_admin.quiz_submissions.quiz_id","lms_admin.quizzes.id",                 "App layer",  "Submission → Quiz"),
        ("lms_admin.certificates.course_id", "lms_admin.courses.id",                  "App layer",  "Certificate → Course"),
        ("lms_admin.coupons.user_id",        "lms_admin.users.id",                    "App layer",  "Coupon creator"),
    ]
)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
#  6. QUICK-REFERENCE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "6. Quick-Reference Summary", 1)

heading(doc, "6.1  All Databases", 2)
schema_table(doc,
    [("Database", 1.4), ("Tables", 0.7), ("Table Names", 5.1)],
    [
        ("lucy_devdb", "9",
         "users · roles · colleges · branches · organisations · "
         "assessments · questionsets · questions · pre_assessment_registrations"),
        ("lms_admin",  "17",
         "users · categories · courses · sections · lessons · quizzes · questions · "
         "quiz_submissions · lesson_completions · lesson_watch_progress · user_progress · "
         "pre_assessment_results · coupons · certificates · live_classes · seo_fields · settings"),
        ("lms_course", "2",
         "courses · enrollments"),
    ]
)

heading(doc, "6.2  Grand Total", 2)
schema_table(doc,
    [("Metric", 2.5), ("Value", 4.7)],
    [
        ("Total Databases",            "3"),
        ("Total Core Tables",          "28  (across 3 databases)"),
        ("lucy_devdb tables",          "9"),
        ("lms_admin tables",           "17"),
        ("lms_course tables",          "2"),
        ("DB Engine",                  "MySQL 8.x on AWS RDS"),
        ("ORM",                        "Sequelize 6  (Node.js)"),
        ("Primary Key styles",         "STRING (UUID-like) in auth/assessment/college/course services; INTEGER auto-increment in admin service"),
        ("Timestamp columns",          "createdAt / updatedAt  (or created_at / updated_at)  on all tables"),
        ("Cross-DB FK enforcement",    "Application layer  (no MySQL-level cross-DB constraints)"),
    ]
)

doc.add_paragraph()
para(doc,
     "This document was auto-generated from the source code of the LMS microservices platform. "
     "All schema details reflect the Sequelize model definitions found in each service's models directory.",
     italic=True, color=RGBColor(0x55, 0x55, 0x55), size=9)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = r"c:\Users\Admin\Downloads\yagnatech1\yagnatech1\lms\LMS_Database_Schema.docx"
doc.save(output_path)
print(f"Document saved: {output_path}")

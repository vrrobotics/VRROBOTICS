"""
One-shot generator for docs/YagnaTech-Database-Schemas.docx.

Reads nothing dynamic — the model definitions are flattened into Python data
structures below so the output is deterministic and reviewable in a diff.
Re-run with `python docs/generate_schema_doc.py` whenever a schema changes.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# --- Schema data ------------------------------------------------------------

# Each table: (name, db_hint, description, columns)
#   columns: list of (field, type, constraints, description)
# `constraints` is a short string: "PK", "PK, AI", "FK -> table.col", "Unique",
# "Not Null", "Default <x>", "Index", combinations comma-separated.

SERVICES = [
    {
        "name": "admin-service",
        "db": "lms_admin (MySQL via Sequelize)",
        "summary": (
            "Core LMS catalog and learning state. Owns courses, sections, lessons, "
            "quizzes, programs, batches, certificates and per-user progress. Two "
            "tables (colleges via the shared College model) actually bind to the "
            "auth-service DB (lucy_devdb) since colleges are canonical there."
        ),
        "tables": [
            {
                "name": "users",
                "model": "User",
                "desc": "Admin/instructor accounts that author courses inside admin-service. Student accounts live in auth-service.users instead.",
                "cols": [
                    ("id", "INTEGER", "PK, AI", "Surrogate primary key."),
                    ("role", "STRING(100)", "Not Null", "Role label (admin / instructor / etc.)."),
                    ("email", "STRING(255)", "Unique, Not Null", "Login email."),
                    ("status", "INTEGER", "", "Account status flag."),
                    ("name", "STRING(255)", "", "Display name."),
                    ("phone", "STRING(255)", "", ""),
                    ("website", "STRING(255)", "", ""),
                    ("skills", "TEXT", "", ""),
                    ("facebook", "TEXT", "", ""),
                    ("twitter", "STRING(255)", "", ""),
                    ("linkedin", "STRING(255)", "", ""),
                    ("address", "STRING(255)", "", ""),
                    ("college_name", "STRING(255)", "", ""),
                    ("college_id", "STRING(255)", "Soft FK -> colleges.clgId", "College admins are scoped via this id; null for root admins."),
                    ("about", "TEXT", "", ""),
                    ("biography", "LONGTEXT", "", ""),
                    ("educations", "LONGTEXT", "", ""),
                    ("photo", "STRING(255)", "", ""),
                    ("email_verified_at", "DATETIME", "", ""),
                    ("password", "STRING(255)", "", "Hashed password."),
                    ("remember_token", "STRING(100)", "", ""),
                    ("paymentkeys", "LONGTEXT", "", ""),
                    ("video_url", "STRING(255)", "", ""),
                    ("created_at", "DATETIME", "", "Sequelize timestamps."),
                    ("updated_at", "DATETIME", "", ""),
                ],
            },
            {
                "name": "categories",
                "model": "Category",
                "desc": "Course taxonomy. Self-referential through parent_id (parent/child categories).",
                "cols": [
                    ("id", "INTEGER", "PK, AI", ""),
                    ("parent_id", "INTEGER", "FK -> categories.id (self)", "Null for top-level categories."),
                    ("title", "STRING(255)", "Not Null", ""),
                    ("slug", "STRING(255)", "", ""),
                    ("icon", "STRING(255)", "", ""),
                    ("sort", "INTEGER", "Default 0", ""),
                    ("status", "STRING(50)", "Default 'active'", ""),
                    ("keywords", "STRING(400)", "", ""),
                    ("description", "STRING(500)", "", ""),
                    ("thumbnail", "STRING(255)", "", ""),
                    ("category_logo", "STRING(255)", "", ""),
                    ("clg_ids", "JSON", "Default []", "Array of clgId strings the category is offered at."),
                    ("created_at", "DATETIME", "", ""),
                    ("updated_at", "DATETIME", "", ""),
                ],
            },
            {
                "name": "courses",
                "model": "Course",
                "desc": "A single course (catalog row). Linked to a category and an admin/instructor creator; scoped to colleges + batches via JSON arrays.",
                "cols": [
                    ("id", "INTEGER", "PK, AI", ""),
                    ("title", "STRING(255)", "Not Null", ""),
                    ("slug", "STRING(255)", "", ""),
                    ("short_description", "TEXT", "", ""),
                    ("user_id", "INTEGER", "FK -> users.id", "Creator (admin)."),
                    ("category_id", "INTEGER", "FK -> categories.id", ""),
                    ("course_type", "STRING(50)", "Default 'general'", ""),
                    ("status", "STRING(50)", "Default 'active'", ""),
                    ("level", "STRING(50)", "", "beginner / intermediate / advanced."),
                    ("language", "STRING(50)", "", ""),
                    ("is_paid", "BOOLEAN", "Default false", ""),
                    ("is_best", "BOOLEAN", "Default false", ""),
                    ("price", "FLOAT", "Default 0", ""),
                    ("discounted_price", "FLOAT", "Default 0", ""),
                    ("discount_flag", "BOOLEAN", "Default false", ""),
                    ("enable_drip_content", "BOOLEAN", "Default false", ""),
                    ("drip_content_settings", "TEXT", "", ""),
                    ("meta_keywords", "TEXT", "", ""),
                    ("meta_description", "TEXT", "", ""),
                    ("thumbnail", "STRING(255)", "", ""),
                    ("banner", "STRING(255)", "", ""),
                    ("preview", "STRING(255)", "", ""),
                    ("description", "LONGTEXT", "", ""),
                    ("requirements", "TEXT", "", ""),
                    ("outcomes", "TEXT", "", ""),
                    ("faqs", "TEXT", "", ""),
                    ("instructor_ids", "TEXT", "", "JSON-encoded list of auth-service userIds."),
                    ("clg_ids", "JSON", "Default []", "Colleges where this course is offered."),
                    ("batch_ids", "JSON", "Default []", "Batches the course is scoped to."),
                    ("average_rating", "FLOAT", "Default 0", ""),
                    ("expiry_period", "INTEGER", "", "Months of access after enrolment; null = lifetime."),
                    ("has_certificate", "BOOLEAN", "Default true", "Whether finishing the course issues a certificate."),
                    ("created_at", "DATETIME", "", ""),
                    ("updated_at", "DATETIME", "", ""),
                ],
            },
            {
                "name": "sections",
                "model": "Section",
                "desc": "An ordered grouping of lessons inside a course.",
                "cols": [
                    ("id", "INTEGER", "PK, AI", ""),
                    ("user_id", "INTEGER", "", "Creator."),
                    ("course_id", "INTEGER", "FK -> courses.id", ""),
                    ("title", "STRING(255)", "", ""),
                    ("sort", "INTEGER", "", "Display order."),
                    ("created_at", "DATETIME", "", ""),
                    ("updated_at", "DATETIME", "", ""),
                ],
            },
            {
                "name": "lessons",
                "model": "Lesson",
                "desc": "A single unit of content (video, quiz, text, etc.) inside a section.",
                "cols": [
                    ("id", "INTEGER", "PK, AI", ""),
                    ("title", "STRING(255)", "", ""),
                    ("user_id", "INTEGER", "", ""),
                    ("course_id", "INTEGER", "FK -> courses.id", ""),
                    ("section_id", "INTEGER", "FK -> sections.id", ""),
                    ("lesson_type", "STRING(255)", "", "video-url / system-video / quiz / etc."),
                    ("duration", "STRING(255)", "", "Format HH:MM:SS."),
                    ("total_mark", "INTEGER", "", "Quiz total marks."),
                    ("pass_mark", "INTEGER", "", "Quiz pass mark."),
                    ("retake", "INTEGER", "", "Quiz retake count."),
                    ("lesson_src", "STRING(255)", "", "Video URL or file path."),
                    ("attachment", "LONGTEXT", "", ""),
                    ("attachment_type", "STRING(255)", "", ""),
                    ("video_type", "TEXT", "", ""),
                    ("thumbnail", "STRING(255)", "", ""),
                    ("is_free", "INTEGER", "", "1 = previewable without enrollment."),
                    ("sort", "INTEGER", "", ""),
                    ("description", "LONGTEXT", "", ""),
                    ("summary", "LONGTEXT", "", ""),
                    ("status", "INTEGER", "", ""),
                    ("created_at", "DATETIME", "", ""),
                    ("updated_at", "DATETIME", "", ""),
                ],
            },
            {
                "name": "questions",
                "model": "Question",
                "desc": "Quiz questions belonging to a quiz-type lesson.",
                "cols": [
                    ("id", "INTEGER", "PK, AI", ""),
                    ("quiz_id", "INTEGER", "FK -> lessons.id", "The lesson row holding the quiz."),
                    ("title", "LONGTEXT", "", "Question text."),
                    ("type", "STRING(255)", "", "mcq / true_false / fill_blanks."),
                    ("answer", "MEDIUMTEXT", "", "JSON-encoded correct answer(s)."),
                    ("options", "LONGTEXT", "", "JSON-encoded options array."),
                    ("sort", "INTEGER", "", ""),
                    ("created_at", "DATETIME", "", ""),
                    ("updated_at", "DATETIME", "", ""),
                ],
            },
            {
                "name": "quiz_submissions",
                "model": "QuizSubmission",
                "desc": "One row per quiz attempt by a student.",
                "cols": [
                    ("id", "INTEGER", "PK, AI", ""),
                    ("quiz_id", "INTEGER", "FK -> lessons.id", ""),
                    ("user_id", "INTEGER", "Soft FK -> auth users.userId", "Note: column is INT, clamps large auth userIds."),
                    ("correct_answer", "LONGTEXT", "", ""),
                    ("wrong_answer", "LONGTEXT", "", ""),
                    ("submits", "LONGTEXT", "", "Full submission payload."),
                    ("score", "INTEGER", "", ""),
                    ("total", "INTEGER", "", ""),
                    ("created_at", "DATETIME", "", ""),
                    ("updated_at", "DATETIME", "", ""),
                ],
            },
            {
                "name": "settings",
                "model": "Setting",
                "desc": "Generic admin settings keyed by `type`.",
                "cols": [
                    ("id", "INTEGER", "PK, AI", ""),
                    ("type", "STRING(255)", "", "Setting key."),
                    ("description", "LONGTEXT", "", "Setting value / payload."),
                    ("created_at", "DATETIME", "", ""),
                    ("updated_at", "DATETIME", "", ""),
                ],
            },
            {
                "name": "live_classes",
                "model": "LiveClass",
                "desc": "Scheduled live sessions for a course.",
                "cols": [
                    ("id", "INTEGER", "PK, AI", ""),
                    ("user_id", "BIGINT", "Soft FK -> auth users.userId", "11-digit auth-service userId."),
                    ("course_id", "INTEGER", "FK -> courses.id", ""),
                    ("class_topic", "STRING(255)", "", ""),
                    ("provider", "STRING(255)", "", "Zoom / Meet / etc."),
                    ("class_date_and_time", "DATETIME", "", ""),
                    ("additional_info", "LONGTEXT", "", ""),
                    ("note", "TEXT", "", ""),
                    ("created_at", "DATETIME", "", ""),
                    ("updated_at", "DATETIME", "", ""),
                ],
            },
            {
                "name": "coupons",
                "model": "Coupon",
                "desc": "Discount coupons.",
                "cols": [
                    ("id", "INTEGER", "PK, AI", ""),
                    ("user_id", "INTEGER", "Not Null", "Owner / creator of the coupon."),
                    ("code", "STRING(255)", "Unique, Not Null", ""),
                    ("discount", "INTEGER", "Not Null", "Percent off."),
                    ("expiry", "BIGINT", "Not Null", "Unix epoch."),
                    ("status", "TINYINT", "Default 1", "1 = active, 0 = disabled."),
                    ("created_at", "DATETIME", "", ""),
                    ("updated_at", "DATETIME", "", ""),
                ],
            },
            {
                "name": "certificates",
                "model": "Certificate",
                "desc": "Issued completion certificates. One row per (user, course).",
                "cols": [
                    ("id", "INTEGER", "PK, AI", ""),
                    ("user_id", "STRING(255)", "Soft FK -> auth users.userId", ""),
                    ("course_id", "INTEGER", "FK -> courses.id", ""),
                    ("identifier", "STRING(100)", "Unique", "Public certificate identifier."),
                    ("title", "STRING(255)", "", ""),
                    ("description", "TEXT", "", ""),
                    ("template_image", "STRING(255)", "", ""),
                    ("status", "TINYINT", "Default 1", "1 = visible, 0 = hidden."),
                    ("issued_at", "DATETIME", "", ""),
                    ("created_at", "DATETIME", "", ""),
                    ("updated_at", "DATETIME", "", ""),
                ],
            },
            {
                "name": "pre_assessment_results",
                "model": "PreAssessmentResult",
                "desc": "Per-attempt result of the pre-assessment used to gate courses. Indexed by (user_id, program_id).",
                "cols": [
                    ("id", "INTEGER", "PK, AI", ""),
                    ("user_id", "BIGINT", "Soft FK -> auth users.userId, Not Null", ""),
                    ("program_id", "INTEGER", "FK -> programs.id", ""),
                    ("score", "FLOAT", "Not Null, Default 0", ""),
                    ("passed", "BOOLEAN", "Not Null, Default false", ""),
                    ("duration_seconds", "INTEGER", "", "Time spent on the assessment."),
                    ("created_at", "DATETIME", "", ""),
                    ("updated_at", "DATETIME", "", ""),
                ],
                "indexes": ["(user_id, program_id)"],
            },
            {
                "name": "user_progress",
                "model": "UserProgress",
                "desc": "Resume point per (user, program). One row per student/program enrollment.",
                "cols": [
                    ("id", "INTEGER", "PK, AI", ""),
                    ("user_id", "BIGINT", "Soft FK -> auth users.userId, Not Null", ""),
                    ("program_id", "INTEGER", "FK -> programs.id, Not Null", ""),
                    ("course_id", "INTEGER", "FK -> courses.id", ""),
                    ("last_lesson_id", "INTEGER", "FK -> lessons.id", ""),
                    ("enrolled", "BOOLEAN", "Not Null, Default false", ""),
                    ("created_at", "DATETIME", "", ""),
                    ("updated_at", "DATETIME", "", ""),
                ],
                "indexes": ["UNIQUE (user_id, program_id)", "(user_id)"],
            },
            {
                "name": "lesson_completions",
                "model": "LessonCompletion",
                "desc": "One row per (user, lesson) once the lesson is considered complete. Drives course progress percentage.",
                "cols": [
                    ("id", "INTEGER", "PK, AI", ""),
                    ("user_id", "BIGINT", "Soft FK -> auth users.userId, Not Null", ""),
                    ("course_id", "INTEGER", "FK -> courses.id, Not Null", ""),
                    ("lesson_id", "INTEGER", "FK -> lessons.id, Not Null", ""),
                    ("created_at", "DATETIME", "", ""),
                    ("updated_at", "DATETIME", "", ""),
                ],
                "indexes": ["UNIQUE (user_id, lesson_id)", "(user_id, course_id)"],
            },
            {
                "name": "lesson_watch_progress",
                "model": "LessonWatchProgress",
                "desc": "Highest playback position per (user, lesson). Used to compute the 30% auto-complete and to resume video playback.",
                "cols": [
                    ("id", "INTEGER", "PK, AI", ""),
                    ("user_id", "BIGINT", "Soft FK -> auth users.userId, Not Null", ""),
                    ("course_id", "INTEGER", "FK -> courses.id, Not Null", ""),
                    ("lesson_id", "INTEGER", "FK -> lessons.id, Not Null", ""),
                    ("current_duration", "INTEGER", "Not Null, Default 0", "Highest playback second reached."),
                    ("created_at", "DATETIME", "", ""),
                    ("updated_at", "DATETIME", "", ""),
                ],
                "indexes": ["UNIQUE (user_id, lesson_id)", "(user_id, course_id)"],
            },
            {
                "name": "languages",
                "model": "Language",
                "desc": "Supported display languages.",
                "cols": [
                    ("id", "INTEGER", "PK, AI", ""),
                    ("name", "STRING(100)", "Unique, Not Null", ""),
                    ("code", "STRING(20)", "", "ISO code."),
                    ("direction", "ENUM('ltr','rtl')", "Not Null, Default 'ltr'", ""),
                    ("is_default", "BOOLEAN", "Not Null, Default false", ""),
                    ("created_at", "DATETIME", "", ""),
                    ("updated_at", "DATETIME", "", ""),
                ],
            },
            {
                "name": "forums",
                "model": "Forum",
                "desc": "Course Q&A threads and replies. Self-referential via parent_id (null = root question).",
                "cols": [
                    ("id", "INTEGER", "PK, AI", ""),
                    ("course_id", "INTEGER", "FK -> courses.id", ""),
                    ("user_id", "BIGINT", "Soft FK -> auth users.userId", ""),
                    ("parent_id", "INTEGER", "FK -> forums.id (self)", "Null on root questions."),
                    ("title", "STRING(255)", "", "'reply' for replies; real title for root."),
                    ("description", "LONGTEXT", "", ""),
                    ("likes", "LONGTEXT", "", "JSON array of liking userIds."),
                    ("dislikes", "LONGTEXT", "", "JSON array of disliking userIds."),
                    ("created_at", "DATETIME", "", ""),
                    ("updated_at", "DATETIME", "", ""),
                ],
            },
            {
                "name": "forum_reports",
                "model": "ForumReport",
                "desc": "Inappropriate-content reports. One row per (post, reporter).",
                "cols": [
                    ("id", "INTEGER", "PK, AI", ""),
                    ("forum_id", "INTEGER", "FK -> forums.id, Not Null", ""),
                    ("user_id", "BIGINT", "Soft FK -> auth users.userId, Not Null", "Reporter."),
                    ("reason", "TEXT", "", ""),
                    ("created_at", "DATETIME", "", ""),
                    ("updated_at", "DATETIME", "", ""),
                ],
                "indexes": ["UNIQUE (forum_id, user_id)"],
            },
            {
                "name": "programs",
                "model": "Program",
                "desc": "Top-level student offerings (AI Frontier, AI Frontier Plus, etc.). Each program bundles courses + colleges + batches.",
                "cols": [
                    ("id", "INTEGER", "PK, AI", ""),
                    ("title", "STRING(255)", "Not Null", ""),
                    ("tagline", "STRING(500)", "", ""),
                    ("icon", "STRING(64)", "Default 'Globe2'", "Lucide icon key."),
                    ("features", "JSON", "Not Null, Default []", "Bullet feature list."),
                    ("sort", "INTEGER", "Default 0", ""),
                    ("is_active", "BOOLEAN", "Default true", ""),
                    ("clg_ids", "JSON", "Default []", "Colleges the program is offered at."),
                    ("course_id", "INTEGER", "Legacy FK -> courses.id", "Single-course legacy column."),
                    ("course_ids", "JSON", "Default []", "Bundled course ids."),
                    ("batch_ids", "JSON", "Default []", "Bundled batch ids."),
                    ("created_at", "DATETIME", "", ""),
                    ("updated_at", "DATETIME", "", ""),
                ],
            },
            {
                "name": "batches",
                "model": "Batch",
                "desc": "Named student cohort at a single college.",
                "cols": [
                    ("id", "INTEGER", "PK, AI", ""),
                    ("clg_id", "STRING(64)", "Soft FK -> colleges.clgId, Not Null", ""),
                    ("name", "STRING(160)", "Not Null", ""),
                    ("description", "STRING(500)", "", ""),
                    ("start_date", "DATE", "", ""),
                    ("end_date", "DATE", "", ""),
                    ("is_active", "BOOLEAN", "Default true", ""),
                    ("created_at", "DATETIME", "", ""),
                    ("updated_at", "DATETIME", "", ""),
                ],
                "indexes": ["(clg_id)"],
            },
            {
                "name": "batch_members",
                "model": "BatchMember",
                "desc": "Link table joining a batch to a student. (batch_id, user_id) is unique.",
                "cols": [
                    ("id", "INTEGER", "PK, AI", ""),
                    ("batch_id", "INTEGER", "FK -> batches.id, Not Null", ""),
                    ("user_id", "STRING(32)", "Soft FK -> auth users.userId, Not Null", ""),
                    ("created_at", "DATETIME", "", ""),
                    ("updated_at", "DATETIME", "", ""),
                ],
                "indexes": ["UNIQUE (batch_id, user_id)", "(user_id)"],
            },
            {
                "name": "email_jobs",
                "model": "EmailJob",
                "desc": "Durable outbound email queue with retry/backoff state.",
                "cols": [
                    ("id", "INTEGER", "PK, AI", ""),
                    ("to_email", "STRING(255)", "Not Null", ""),
                    ("subject", "STRING(255)", "Not Null", ""),
                    ("html", "LONGTEXT", "Not Null", ""),
                    ("batch_id", "INTEGER", "Soft FK -> batches.id", ""),
                    ("user_id", "STRING(32)", "Soft FK -> auth users.userId", ""),
                    ("status", "ENUM(queued, sending, sent, failed)", "Not Null, Default 'queued'", ""),
                    ("attempts", "INTEGER", "Not Null, Default 0", ""),
                    ("last_error", "STRING(500)", "", ""),
                    ("next_attempt_at", "DATETIME", "Not Null, Default NOW", "Backoff schedule."),
                    ("sent_at", "DATETIME", "", ""),
                    ("failed_at", "DATETIME", "", ""),
                    ("created_at", "DATETIME", "", ""),
                    ("updated_at", "DATETIME", "", ""),
                ],
                "indexes": ["(status, next_attempt_at)", "(batch_id)", "(user_id)"],
            },
            {
                "name": "seo_fields",
                "model": "SeoField",
                "desc": "Per-course or per-route SEO metadata.",
                "cols": [
                    ("id", "INTEGER", "PK, AI", ""),
                    ("course_id", "INTEGER", "FK -> courses.id", ""),
                    ("route", "STRING(255)", "", ""),
                    ("name_route", "STRING(255)", "", ""),
                    ("meta_title", "STRING(255)", "", ""),
                    ("meta_description", "TEXT", "", ""),
                    ("meta_keywords", "TEXT", "", ""),
                    ("meta_robot", "STRING(100)", "", ""),
                    ("canonical_url", "STRING(255)", "", ""),
                    ("custom_url", "STRING(255)", "", ""),
                    ("json_ld", "TEXT", "", ""),
                    ("og_title", "STRING(255)", "", ""),
                    ("og_description", "TEXT", "", ""),
                    ("og_image", "STRING(255)", "", ""),
                    ("created_at", "DATETIME", "", ""),
                    ("updated_at", "DATETIME", "", ""),
                ],
            },
        ],
    },

    {
        "name": "auth-service",
        "db": "lucy_devdb (MySQL via Sequelize)",
        "summary": (
            "Owns the canonical user identity (students, instructors, admins) "
            "and role catalog. Other services treat users.userId as a soft FK."
        ),
        "tables": [
            {
                "name": "roles",
                "model": "Role",
                "desc": "Role catalog.",
                "cols": [
                    ("roleId", "STRING", "PK", ""),
                    ("role", "ENUM(student, instructor, admin, auditor)", "Not Null", ""),
                ],
            },
            {
                "name": "users",
                "model": "User",
                "desc": "Canonical user identity. userId is the auth-service primary key referenced (without DB FK) by every other service.",
                "cols": [
                    ("userId", "STRING", "PK", "11-digit numeric string."),
                    ("name", "STRING", "Not Null", ""),
                    ("email", "STRING", "Unique, Not Null", ""),
                    ("passwordHash", "STRING", "Not Null", ""),
                    ("phone", "STRING", "", ""),
                    ("dob", "DATE", "", ""),
                    ("gender", "ENUM(male, female)", "", ""),
                    ("yearOfEducation", "STRING", "", ""),
                    ("branchId", "STRING", "Soft FK -> branches.branchId", ""),
                    ("collegeId", "STRING", "Soft FK -> colleges.clgId", ""),
                    ("yearOfStudy", "INTEGER", "", ""),
                    ("educationLevel", "ENUM(inter, bachelor, master, phd, other)", "", ""),
                    ("branch", "STRING", "", "Free-text academic branch."),
                    ("collegeName", "STRING", "", ""),
                    ("graduationYear", "STRING", "", ""),
                    ("collegeCode", "STRING", "", ""),
                    ("orgId", "STRING", "Soft FK -> organisations.orgId", ""),
                    ("assessmentId", "STRING", "Soft FK -> assessments.assessmentId", ""),
                    ("programInterested", "STRING", "", ""),
                    ("expertise", "STRING(255)", "", "Instructor field."),
                    ("bio", "STRING(1000)", "", "Instructor field."),
                    ("yearsOfExperience", "INTEGER", "", "Instructor field."),
                    ("linkedinUrl", "STRING(255)", "", "Instructor field."),
                    ("profileStatus", "ENUM(active, inactive, pending)", "Default 'pending'", ""),
                    ("location", "STRING", "", ""),
                    ("address", "STRING(255)", "", ""),
                    ("lastLogin", "DATETIME", "", ""),
                    ("preScore", "INTEGER", "", "Pre-assessment score."),
                    ("preScoreDuration", "INTEGER", "", "Seconds spent on the pre-assessment."),
                    ("postScore", "INTEGER", "", ""),
                    ("postScoreDuration", "INTEGER", "", "Seconds spent on the post-assessment."),
                    ("refreshToken", "STRING(1024)", "", ""),
                    ("roleId", "STRING", "FK -> roles.roleId, Not Null", ""),
                    ("createdAt", "DATETIME", "", ""),
                    ("updatedAt", "DATETIME", "", ""),
                ],
            },
        ],
    },

    {
        "name": "assessment-service",
        "db": "Sequelize (dedicated assessment DB)",
        "summary": (
            "Pre/post assessment question banks, assessment instances, and the "
            "pre-assessment registration / declaration flow."
        ),
        "tables": [
            {
                "name": "questions",
                "model": "Question",
                "desc": "Reusable assessment question bank.",
                "cols": [
                    ("quesId", "STRING", "PK", ""),
                    ("question", "TEXT", "Not Null", ""),
                    ("correctAns", "STRING", "Not Null", ""),
                    ("options", "JSON", "Not Null", "Shape: {option1, option2, option3, option4}."),
                    ("category", "STRING", "", ""),
                    ("questionSeverity", "ENUM(easy, medium, hard)", "", ""),
                    ("createdAt", "DATETIME", "", ""),
                    ("updatedAt", "DATETIME", "", ""),
                ],
            },
            {
                "name": "questionsets",
                "model": "QuestionSet",
                "desc": "Named bundle of questions referenced by an assessment.",
                "cols": [
                    ("setId", "STRING", "PK", ""),
                    ("setName", "STRING", "Not Null", ""),
                    ("category", "STRING", "", ""),
                    ("questions", "JSON", "Not Null, Default []", "Array of quesIds."),
                    ("createdAt", "DATETIME", "", ""),
                    ("updatedAt", "DATETIME", "", ""),
                ],
            },
            {
                "name": "assessments",
                "model": "Assessment",
                "desc": "An assessment instance (pre or post) bound to a question set.",
                "cols": [
                    ("assessmentId", "STRING", "PK", ""),
                    ("type", "ENUM(pre, post)", "Not Null", ""),
                    ("setId", "STRING", "FK -> questionsets.setId, Not Null", ""),
                    ("startAt", "DATETIME", "", ""),
                    ("endAt", "DATETIME", "", ""),
                    ("score", "FLOAT", "", ""),
                    ("timer", "INTEGER", "", "Duration in seconds/minutes."),
                    ("status", "ENUM(not-started, available, in-progress, completed, expired)", "Not Null, Default 'not-started'", ""),
                    ("clgIds", "JSON", "Not Null, Default []", "Colleges offered to."),
                    ("courseIds", "JSON", "Not Null, Default []", "Courses bundled with."),
                    ("createdAt", "DATETIME", "", ""),
                    ("updatedAt", "DATETIME", "", ""),
                ],
            },
            {
                "name": "pre_assessment_registrations",
                "model": "PreAssessmentRegistration",
                "desc": "Onboarding registration captured before a student attempts the pre-assessment.",
                "cols": [
                    ("registrationId", "STRING", "PK", ""),
                    ("userId", "STRING", "Soft FK -> auth users.userId", "Null for walk-ins."),
                    ("fullName", "STRING(120)", "Not Null", ""),
                    ("email", "STRING(160)", "Not Null", "Email format validated."),
                    ("phoneNumber", "STRING(20)", "Not Null", ""),
                    ("gender", "ENUM(...GENDERS)", "Not Null", ""),
                    ("selectedProgramId", "INTEGER", "Soft FK -> programs.id", ""),
                    ("selectedProgram", "STRING(255)", "Not Null", "Frozen program title snapshot."),
                    ("uploadedCollegeProof", "JSON", "Not Null", "{fileName, originalName, mimeType, size, url, storedAt}."),
                    ("declarationAccepted", "BOOLEAN", "Not Null, Default false", "Must be true."),
                    ("assessmentStatus", "ENUM(...ASSESSMENT_STATUSES)", "Not Null, Default 'registered'", ""),
                    ("assessmentStartedAt", "DATETIME", "", ""),
                    ("submittedFromIp", "STRING(64)", "", ""),
                    ("submittedUserAgent", "STRING(255)", "", ""),
                    ("createdAt", "DATETIME", "", ""),
                    ("updatedAt", "DATETIME", "", ""),
                ],
                "indexes": ["(userId)", "(email)", "(selectedProgram)", "(assessmentStatus)"],
            },
        ],
    },

    {
        "name": "college-service",
        "db": "Sequelize (shared with auth-service)",
        "summary": (
            "Owns the canonical College and Branch catalog. College rows are "
            "shared with auth-service via the same DB, and admin-service binds "
            "its College model to this DB for read access."
        ),
        "tables": [
            {
                "name": "colleges",
                "model": "College",
                "desc": "A college / institution. clgId is the canonical id referenced (without DB FK) across services.",
                "cols": [
                    ("clgId", "STRING", "PK", ""),
                    ("accesskey", "STRING", "Unique, Not Null", "Per-college signup access key."),
                    ("clgName", "STRING", "Not Null", ""),
                    ("clgAddress", "TEXT", "", ""),
                    ("orgId", "STRING", "FK -> organisations.orgId", ""),
                    ("branchIds", "JSON", "", "Array of branch PKs (e.g. ['CSE','ECE'])."),
                    ("isActive", "BOOLEAN", "Not Null, Default true", "Toggled from Manage Colleges (admin-service model adds this column)."),
                    ("createdAt", "DATETIME", "", ""),
                    ("updatedAt", "DATETIME", "", ""),
                ],
            },
            {
                "name": "branches",
                "model": "Branch",
                "desc": "Academic branch catalog (CSE, ECE, etc.). No timestamps.",
                "cols": [
                    ("branchId", "STRING", "PK", ""),
                    ("branchName", "STRING", "Unique, Not Null", ""),
                ],
            },
        ],
    },

    {
        "name": "course-service",
        "db": "Sequelize (dedicated course DB)",
        "summary": (
            "A lighter, JSON-modeled course catalog and enrollment store used "
            "for the public catalog flow. Separate from the richer "
            "admin-service.courses table (which powers the admin authoring UI)."
        ),
        "tables": [
            {
                "name": "courses",
                "model": "Course",
                "desc": "Catalog course row with modules stored as JSON. Distinct from admin-service.courses.",
                "cols": [
                    ("courseId", "STRING", "PK", ""),
                    ("title", "STRING", "Not Null", ""),
                    ("description", "TEXT", "", ""),
                    ("duration", "INTEGER", "Not Null", "Hours or minutes."),
                    ("isPreAssessmentNeeded", "BOOLEAN", "Default false", ""),
                    ("modules", "JSON", "", "Array of {moduleId, title, duration}."),
                    ("clgIds", "JSON", "Not Null, Default []", "Colleges the course is offered at."),
                    ("instructorId", "STRING", "Soft FK -> auth users.userId", "Instructor account id."),
                    ("createdAt", "DATETIME", "", ""),
                    ("updatedAt", "DATETIME", "", ""),
                ],
            },
            {
                "name": "enrollments",
                "model": "Enrollment",
                "desc": "Student enrollment in a course.",
                "cols": [
                    ("enrollmentId", "STRING", "PK", ""),
                    ("userId", "STRING", "Soft FK -> auth users.userId, Not Null", ""),
                    ("courseId", "STRING", "FK -> courses.courseId, Not Null", ""),
                    ("status", "ENUM(enrolled, in-progress, completed, dropped)", "Not Null, Default 'enrolled'", ""),
                    ("enrolledAt", "DATETIME", "Default NOW", ""),
                    ("completedAt", "DATETIME", "", ""),
                    ("createdAt", "DATETIME", "", ""),
                    ("updatedAt", "DATETIME", "", ""),
                ],
            },
        ],
    },

    {
        "name": "organization-service",
        "db": "Sequelize (dedicated organisation DB)",
        "summary": "Owns organisations (parent of colleges).",
        "tables": [
            {
                "name": "organisations",
                "model": "Organisation",
                "desc": "Top-level organisation. Referenced by colleges.orgId.",
                "cols": [
                    ("orgId", "STRING", "PK", ""),
                    ("accesskey", "STRING", "Unique, Not Null", ""),
                    ("orgName", "STRING", "Not Null", ""),
                    ("orgState", "STRING", "", ""),
                    ("orgCountry", "STRING", "", ""),
                    ("orgAddress", "TEXT", "", ""),
                    ("orgPin", "STRING", "", "Postal code (string to support non-numeric formats)."),
                    ("createdAt", "DATETIME", "", ""),
                    ("updatedAt", "DATETIME", "", ""),
                ],
            },
        ],
    },

    {
        "name": "payment-service",
        "db": "Sequelize (shared schema with college-service)",
        "summary": (
            "Re-declares the College and Branch models locally so payment-service "
            "can read them without a cross-service call. Same row shapes as "
            "college-service."
        ),
        "tables": [
            {
                "name": "colleges",
                "model": "College",
                "desc": "Local view of the canonical colleges table.",
                "cols": [
                    ("clgId", "STRING", "PK", ""),
                    ("accesskey", "STRING", "Unique, Not Null", ""),
                    ("clgName", "STRING", "Not Null", ""),
                    ("clgAddress", "TEXT", "", ""),
                    ("orgId", "STRING", "FK -> organisations.orgId", ""),
                    ("branchIds", "JSON", "", ""),
                    ("createdAt", "DATETIME", "", ""),
                    ("updatedAt", "DATETIME", "", ""),
                ],
            },
            {
                "name": "branches",
                "model": "Branch",
                "desc": "Local view of the canonical branches table.",
                "cols": [
                    ("branchId", "STRING", "PK", ""),
                    ("branchName", "STRING", "Unique, Not Null", ""),
                ],
            },
        ],
    },
]


# --- Doc rendering ----------------------------------------------------------

BRAND_TEAL = RGBColor(0x17, 0x73, 0x85)
HEADER_FILL = "177385"
ALT_ROW_FILL = "F4F8F9"


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "BFD3D8")
        tcBorders.append(border)
    tc_pr.append(tcBorders)


def add_table(doc, columns_def, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(columns_def))
    table.autofit = False
    table.allow_autofit = False

    # Header row
    header = table.rows[0]
    for i, (label, width_cm) in enumerate(columns_def):
        cell = header.cells[i]
        cell.width = Cm(width_cm)
        set_cell_bg(cell, HEADER_FILL)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        run.font.name = "Calibri"

    # Body rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, value in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.width = Cm(columns_def[c_idx][1])
            set_cell_borders(cell)
            if r_idx % 2 == 1:
                set_cell_bg(cell, ALT_ROW_FILL)
            p = cell.paragraphs[0]
            run = p.add_run(value or "")
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"
            if c_idx == 0:
                run.bold = True
            if c_idx == 2 and value and "PK" in value:
                run.font.color.rgb = BRAND_TEAL
                run.bold = True

    return table


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BRAND_TEAL
        run.font.name = "Calibri"
    return h


def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("YagnaTech Foundation")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = BRAND_TEAL
    run.font.name = "Calibri"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("Backend Database Schemas")
    sub_run.font.size = Pt(18)
    sub_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    sub_run.font.name = "Calibri"

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run("Reference document — auto-generated from Sequelize models")
    meta_run.italic = True
    meta_run.font.size = Pt(11)
    meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # Legend / overview
    add_heading(doc, "Overview", level=1)
    overview = doc.add_paragraph()
    overview.add_run(
        "Each microservice owns its own Sequelize models. Some services share a "
        "physical database (auth-service and college-service share lucy_devdb), "
        "and a few models are re-declared across services so each one can read "
        "the same row shape without an inter-service call. Foreign keys marked "
        "as \"Soft FK\" are enforced at the application layer only (no DB-level "
        "constraint), typically because the referenced row lives in a different "
        "service's database."
    )

    add_heading(doc, "Column-constraint legend", level=2)
    legend_rows = [
        ("PK", "Primary key"),
        ("AI", "AUTO_INCREMENT"),
        ("FK -> table.col", "DB-level foreign key"),
        ("Soft FK -> table.col", "Application-level reference (no DB constraint)"),
        ("Unique", "Unique index on the column"),
        ("Not Null", "NOT NULL constraint"),
        ("Default <x>", "Column default value"),
    ]
    add_table(doc, [("Token", 4.5), ("Meaning", 12)], legend_rows)

    doc.add_page_break()

    # Per-service sections
    for svc in SERVICES:
        add_heading(doc, svc["name"], level=1)

        db_p = doc.add_paragraph()
        db_label = db_p.add_run("Database: ")
        db_label.bold = True
        db_p.add_run(svc["db"])

        doc.add_paragraph(svc["summary"])

        for table in svc["tables"]:
            add_heading(doc, f"{table['name']}  ({table['model']})", level=2)
            doc.add_paragraph(table["desc"])

            cols_def = [
                ("Field", 4.5),
                ("Type", 4.5),
                ("Constraints", 4.5),
                ("Description", 5.0),
            ]
            rows = [(c[0], c[1], c[2], c[3]) for c in table["cols"]]
            add_table(doc, cols_def, rows)

            if table.get("indexes"):
                idx_p = doc.add_paragraph()
                idx_p.add_run("Indexes: ").bold = True
                idx_p.add_run("; ".join(table["indexes"]))

            doc.add_paragraph()  # spacer

        doc.add_page_break()

    out_path = Path(__file__).resolve().parent / "YagnaTech-Database-Schemas.docx"
    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
